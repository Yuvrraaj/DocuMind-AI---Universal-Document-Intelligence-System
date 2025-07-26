import os
import json
import time
import re
import threading
import hashlib
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, font
from functools import wraps
from typing import List, Optional, Dict, Any
from dataclasses import dataclass
from pydantic import BaseModel
import logging
from datetime import datetime
from PIL import Image
import io

# Import config
try:
    import config
except ImportError:
    print("Please create config.py with your GEMINI_API_KEY")
    exit(1)

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)

# Configure Gemini
import google.generativeai as genai
genai.configure(api_key=os.environ["GOOGLE_API_KEY"])

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
import PyPDF2
from docx import Document as DocxDocument

# Color scheme
COLORS = {
    'primary': '#2E86AB',      # Blue
    'secondary': '#A23B72',    # Purple
    'success': '#F18F01',      # Orange
    'danger': '#C73E1D',       # Red
    'warning': '#F4B942',      # Yellow
    'info': '#59C9A5',         # Teal
    'light': '#F8F9FA',        # Light gray
    'dark': '#2C3E50',         # Dark blue-gray
    'white': '#FFFFFF',
    'bg_primary': '#E8F4FD',   # Light blue
    'bg_secondary': '#F3E8F7', # Light purple
    'text_dark': '#2C3E50',
    'text_light': '#6C757D',
    'border': '#DEE2E6'
}

# ------------------ Rate limiting and caching ------------------

def rate_limit(calls_per_minute: int):
    """Rate limiting decorator to prevent API quota exceeded errors"""
    min_interval = 60.0 / calls_per_minute
    def decorator(func):
        last_called = [0.0]
        @wraps(func)
        def wrapper(*args, **kwargs):
            elapsed = time.time() - last_called[0]
            wait_time = min_interval - elapsed
            if wait_time > 0:
                logger.debug(f"Rate limiting: sleeping for {wait_time:.2f} seconds")
                time.sleep(wait_time)
            ret = func(*args, **kwargs)
            last_called[0] = time.time()
            return ret
        return wrapper
    return decorator

class APICache:
    """Simple file-based cache for API responses"""
    def __init__(self, cache_dir=config.CACHE_DIR):
        self.cache_dir = cache_dir
        os.makedirs(self.cache_dir, exist_ok=True)

    def _get_cache_key(self, prompt: str) -> str:
        return hashlib.md5(prompt.encode('utf-8')).hexdigest()

    def get(self, prompt: str) -> Optional[str]:
        cache_key = self._get_cache_key(prompt)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.json")
        
        if os.path.exists(cache_file):
            try:
                with open(cache_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    return data.get("content")
            except Exception as e:
                logger.warning(f"Failed to read cache: {e}")
        return None

    def set(self, prompt: str, content: str):
        cache_key = self._get_cache_key(prompt)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.json")
        
        try:
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump({"content": content, "timestamp": time.time()}, f)
        except Exception as e:
            logger.warning(f"Failed to write cache: {e}")

class UsageTracker:
    """Track API usage to monitor quota consumption"""
    def __init__(self):
        self.request_count = 0
        self.daily_requests = 0
        self.last_reset = datetime.now().date()
        self.start_time = time.time()

    def log_request(self):
        current_date = datetime.now().date()
        if current_date != self.last_reset:
            self.daily_requests = 0
            self.last_reset = current_date

        self.request_count += 1
        self.daily_requests += 1

        elapsed_minutes = (time.time() - self.start_time) / 60
        rpm = self.request_count / max(elapsed_minutes, 1)

        logger.info(f"API Usage - Total: {self.request_count}, Daily: {self.daily_requests}, RPM: {rpm:.2f}")

        if rpm > 12:
            logger.warning("⚠️ Approaching rate limit!")
        if self.daily_requests > 1400:
            logger.warning("⚠️ Approaching daily limit!")

        return {
            "total_requests": self.request_count,
            "daily_requests": self.daily_requests,
            "requests_per_minute": rpm
        }

# ------------------ Data Classes ------------------

@dataclass
class QueryStructure:
    age: Optional[int] = None
    gender: Optional[str] = None
    procedure: Optional[str] = None
    location: Optional[str] = None
    policy_duration: Optional[str] = None
    policy_type: Optional[str] = None
    sum_insured: Optional[str] = None
    raw_query: str = ""

class AnalysisResponse(BaseModel):
    answer: str
    justification: str
    confidence_score: float
    key_points: List[str]
    document_references: List[str]
    sources: List[Dict[str, Any]]
    applicable_sections: List[str]

# ------------------ Core System Classes ------------------

class SafeModelClient:
    """Gemini client with rate limiting, caching, and retry logic"""
    def __init__(self, model_name: str = "gemini-1.5-flash"):
        self.model = genai.GenerativeModel(model_name)
        self.cache = APICache()
        self.usage_tracker = UsageTracker()
        
        # Apply rate limiting
        self.rate_limited_generate = rate_limit(config.API_RATE_LIMIT)(self._generate_content_internal)

    def _generate_content_internal(self, prompt: str) -> str:
        """Internal method for actual API call"""
        response = self.model.generate_content(prompt)
        return response.text or ""

    def generate_content(self, prompt: str) -> str:
        """Generate content with caching and retry logic"""
        # Check cache first
        cached_content = self.cache.get(prompt)
        if cached_content:
            logger.info("✅ Cache hit - using cached response")
            return cached_content

        # Track usage
        usage_stats = self.usage_tracker.log_request()

        # Attempt API call with retries
        for attempt in range(config.MAX_RETRIES):
            try:
                logger.info(f"🔄 API call attempt {attempt + 1}")
                content = self.rate_limited_generate(prompt)
                
                # Cache the response
                self.cache.set(prompt, content)
                logger.info("✅ API call successful")
                return content

            except Exception as e:
                error_str = str(e)
                if "429" in error_str or "quota" in error_str.lower():
                    delay = config.RETRY_BASE_DELAY * (2 ** attempt)
                    logger.warning(f"⚠️ Rate limit hit! Retrying in {delay} seconds...")
                    time.sleep(delay)
                elif "400" in error_str:
                    logger.error(f"❌ Bad request error: {e}")
                    raise e
                else:
                    logger.error(f"❌ API error: {e}")
                    if attempt == config.MAX_RETRIES - 1:
                        raise e
                    time.sleep(5)  # Short delay for other errors

        raise RuntimeError("🔥 Exceeded maximum retry attempts for API calls")

    def get_usage_stats(self):
        return self.usage_tracker.log_request()

class EnhancedDocumentProcessor:
    """Enhanced document processor with support for multiple formats"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
        self.model_client = SafeModelClient('gemini-1.5-flash')

    def load_pdf(self, filepath: str) -> str:
        """Extract text from PDF with better error handling"""
        try:
            text = ""
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"\n--- PAGE {page_num + 1} ---\n{page_text.strip()}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                        continue

            if not text.strip():
                raise ValueError("No readable text found in PDF")
            return text.strip()
        except Exception as e:
            raise Exception(f"Error loading PDF: {e}")

    def load_docx(self, filepath: str) -> str:
        """Extract text from Word document"""
        try:
            doc = DocxDocument(filepath)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"

            if not text.strip():
                raise ValueError("No readable text found in Word document")
            return text.strip()
        except Exception as e:
            raise Exception(f"Error loading DOCX: {e}")

    def load_image(self, filepath: str) -> str:
        """Extract text from image using Gemini Vision"""
        try:
            # Load and prepare image
            with open(filepath, 'rb') as f:
                image_data = f.read()

            image = Image.open(io.BytesIO(image_data))
            
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Resize if too large
            max_size = 1024
            if max(image.size) > max_size:
                ratio = max_size / max(image.size)
                new_size = tuple(int(dim * ratio) for dim in image.size)
                image = image.resize(new_size, Image.Resampling.LANCZOS)

            # Convert back to bytes
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='JPEG', quality=85)
            img_byte_arr = img_byte_arr.getvalue()

            prompt = """
            Extract ALL text content from this document image. Focus on:
            1. Policy terms and conditions
            2. Coverage details and benefits
            3. Exclusions and limitations
            4. Waiting periods and eligibility
            5. Premium and sum insured details
            6. Claim procedures and contact information
            
            Return the extracted text with proper structure and formatting.
            """

            # Use vision model for text extraction
            vision_model = genai.GenerativeModel('gemini-1.5-flash')
            response = vision_model.generate_content([prompt, {"mime_type": "image/jpeg", "data": img_byte_arr}])

            if not response.text or not response.text.strip():
                raise Exception("No text could be extracted from the image")

            return response.text.strip()
        except Exception as e:
            raise Exception(f"Error loading image: {e}")

    def load_text(self, filepath: str) -> str:
        """Load plain text file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
            if not text.strip():
                raise ValueError("Text file is empty")
            return text.strip()
        except Exception as e:
            raise Exception(f"Error loading text file: {e}")

    def analyze_document_content(self, text: str, filename: str) -> str:
        """Analyze document content for better understanding"""
        try:
            prompt = f"""
            Analyze this insurance document and provide a structured summary:

            DOCUMENT: {filename}
            CONTENT: {text[:3000]}...

            Provide:
            1. Document type and purpose
            2. Key coverage areas and benefits
            3. Important terms, conditions, and restrictions
            4. Waiting periods and eligibility criteria
            5. Exclusions and limitations
            6. Premium and sum insured information
            7. Claim procedures

            Format as a comprehensive analysis for query processing.
            """

            analysis = self.model_client.generate_content(prompt)
            return f"=== DOCUMENT ANALYSIS ===\n{analysis}\n\n=== ORIGINAL CONTENT ===\n{text}"
        except Exception as e:
            logger.warning(f"Could not analyze document: {e}")
            return text

    def process_document(self, filepath: str) -> List[Document]:
        """Process document and return enhanced chunks"""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        filename = os.path.basename(filepath)
        file_ext = os.path.splitext(filepath)[1].lower()

        logger.info(f"📄 Processing: {filename}")

        # Extract text based on file type
        if file_ext == '.pdf':
            text = self.load_pdf(filepath)
        elif file_ext == '.docx':
            text = self.load_docx(filepath)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            text = self.load_image(filepath)
        elif file_ext in ['.txt', '.md']:
            text = self.load_text(filepath)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Clean text
        text = self._clean_text(text)

        # Analyze content for better understanding
        enhanced_text = self.analyze_document_content(text, filename)

        # Split into chunks
        chunks = self.text_splitter.split_text(enhanced_text)
        
        # Create Document objects
        documents = []
        for i, chunk in enumerate(chunks):
            if chunk.strip():
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": filepath,
                        "filename": filename,
                        "chunk_id": i,
                        "file_type": file_ext,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)

        logger.info(f"✅ Created {len(documents)} chunks for {filename}")
        return documents

    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        # Remove special characters
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\xff]', '', text)
        return text.strip()

class EnhancedQueryProcessor:
    """Enhanced query processor with intelligent document understanding"""

    def __init__(self):
        self.model_client = SafeModelClient('gemini-1.5-flash')
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        self.vector_store: Optional[Chroma] = None
        self.document_summaries = []

    def create_vector_store(self, documents: List[Document], persist_path="./chroma_db"):
        """Create vector store with document analysis"""
        try:
            import shutil
            if os.path.exists(persist_path):
                shutil.rmtree(persist_path)

            self.vector_store = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                persist_directory=persist_path
            )
            self.vector_store.persist()

            # Create document summaries
            self._create_document_summaries(documents)
            
            logger.info(f"🔍 Vector store created with {len(documents)} documents")
        except Exception as e:
            raise Exception(f"Error creating vector store: {e}")

    def _create_document_summaries(self, documents: List[Document]):
        """Create intelligent summaries of loaded documents"""
        try:
            # Group documents by source
            file_groups = {}
            for doc in documents:
                source = doc.metadata.get('source', 'unknown')
                if source not in file_groups:
                    file_groups[source] = []
                file_groups[source].append(doc)

            self.document_summaries = []
            for source, docs in file_groups.items():
                # Use first few chunks for summary
                combined_content = "\n".join([doc.page_content for doc in docs[:2]])

                summary_prompt = f"""
                Create a comprehensive summary of this insurance document:

                FILE: {os.path.basename(source)}
                CONTENT: {combined_content[:2500]}

                Provide:
                1. Document type and main purpose
                2. Coverage areas and benefits
                3. Key terms and conditions
                4. Important restrictions or exclusions
                5. Waiting periods and eligibility
                6. Claim procedures and requirements

                Make this summary useful for answering user queries about the document.
                """

                try:
                    summary = self.model_client.generate_content(summary_prompt)
                    self.document_summaries.append({
                        'source': source,
                        'filename': os.path.basename(source),
                        'summary': summary,
                        'chunk_count': len(docs)
                    })
                except Exception as e:
                    logger.warning(f"Could not create summary for {source}: {e}")

        except Exception as e:
            logger.warning(f"Error creating summaries: {e}")

    def search_relevant_content(self, query: str, k: int = 6) -> List[Document]:
        """Search for relevant content with enhanced query"""
        if not self.vector_store:
            raise ValueError("Vector store not initialized")

        # Enhance query with insurance context
        enhanced_query = f"{query} insurance policy coverage benefits exclusions claims terms conditions"
        
        relevant_docs = self.vector_store.similarity_search(enhanced_query, k=k)
        logger.info(f"🔍 Found {len(relevant_docs)} relevant content pieces")
        return relevant_docs

    def process_query(self, query: str) -> AnalysisResponse:
        """Process query and return comprehensive analysis"""
        if not self.vector_store:
            raise ValueError("No documents loaded. Please upload and process documents first.")

        logger.info(f"🤔 Processing query: {query}")

        # Search for relevant content
        relevant_docs = self.search_relevant_content(query, k=6)

        if not relevant_docs:
            return AnalysisResponse(
                answer="No relevant information found in the uploaded documents.",
                justification="Could not find relevant content to answer your query.",
                confidence_score=0.0,
                key_points=[],
                document_references=[],
                sources=[],
                applicable_sections=[]
            )

        # Prepare context from relevant documents
        context = self._prepare_context(relevant_docs[:config.MAX_CONTEXT_DOCS])

        # Get document summaries for context
        summaries_text = self._get_summaries_text()

        # Create comprehensive analysis prompt
        analysis_prompt = f"""
        You are an expert insurance document analyst. Answer the user's query based on the uploaded documents.

        USER QUERY: {query}

        RELEVANT DOCUMENT EXCERPTS:
        {context}

        DOCUMENT SUMMARIES:
        {summaries_text}

        Provide a comprehensive analysis in JSON format:
        {{
            "answer": "Direct, detailed answer to the user's query",
            "justification": "Detailed explanation with specific document references and quotes",
            "confidence_score": 0.0 to 1.0,
            "key_points": ["Important point 1", "Important point 2", "Important point 3"],
            "document_references": ["Specific quote 1", "Specific quote 2"],
            "applicable_sections": ["Section name 1", "Section name 2"]
        }}

        Be specific, reference exact content from documents, and provide actionable information.
        Return ONLY the JSON response.
        """

        try:
            response_text = self.model_client.generate_content(analysis_prompt)

            # Clean and parse JSON response
            response_text = self._clean_json_response(response_text)
            result_data = json.loads(response_text)

            # Add source information
            sources = []
            for i, doc in enumerate(relevant_docs):
                sources.append({
                    'source_id': f"source_{i+1}",
                    'filename': doc.metadata.get('filename', 'unknown'),
                    'content_preview': doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                    'relevance': max(0.1, 1.0 - (i * 0.15))
                })

            return AnalysisResponse(
                answer=result_data.get('answer', 'No answer provided'),
                justification=result_data.get('justification', 'No justification provided'),
                confidence_score=min(1.0, max(0.0, result_data.get('confidence_score', 0.5))),
                key_points=result_data.get('key_points', []),
                document_references=result_data.get('document_references', []),
                sources=sources,
                applicable_sections=result_data.get('applicable_sections', [])
            )

        except Exception as e:
            logger.error(f"Error processing query: {e}")
            return AnalysisResponse(
                answer=f"Error processing query: {str(e)}",
                justification="An error occurred during analysis",
                confidence_score=0.0,
                key_points=[],
                document_references=[],
                sources=[],
                applicable_sections=[]
            )

    def _prepare_context(self, docs: List[Document]) -> str:
        """Prepare context from relevant documents"""
        context_parts = []
        for i, doc in enumerate(docs, 1):
            filename = doc.metadata.get('filename', 'unknown')
            content = doc.page_content[:1200] + "..." if len(doc.page_content) > 1200 else doc.page_content
            context_parts.append(f"--- SOURCE {i}: {filename} ---\n{content}\n")
        return "\n".join(context_parts)

    def _get_summaries_text(self) -> str:
        """Get formatted document summaries"""
        if not self.document_summaries:
            return "No document summaries available."
        
        summaries = []
        for summary in self.document_summaries:
            summaries.append(f"FILE: {summary['filename']}\nSUMMARY: {summary['summary'][:800]}...\n")
        return "\n".join(summaries)

    def _clean_json_response(self, response_text: str) -> str:
        """Clean JSON response from markdown formatting"""
        if response_text.startswith('```'):
            response_text = response_text[7:-3].strip()
        elif response_text.startswith('```'):
            response_text = response_text[3:-3].strip()
        
        # Find JSON boundaries
        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1
        if start_idx != -1 and end_idx > start_idx:
            response_text = response_text[start_idx:end_idx]
        
        return response_text

# ------------------ Scrollable Frame Widget ------------------

class ScrollableFrame(tk.Frame):
    """A scrollable frame widget"""
    def __init__(self, parent, bg=None, **kwargs):
        super().__init__(parent, **kwargs)
        
        # Create canvas and scrollbar
        self.canvas = tk.Canvas(self, bg=bg or COLORS['light'], highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas, bg=bg or COLORS['light'])
        
        # Configure scrolling
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        # Pack canvas and scrollbar
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        
        # Bind mouse wheel events
        self.bind_mousewheel()
        
        # Configure canvas window width
        self.canvas.bind('<Configure>', self._on_canvas_configure)
    
    def _on_canvas_configure(self, event):
        """Configure canvas window width"""
        canvas_width = event.width
        self.canvas.itemconfig(self.canvas_window, width=canvas_width)
    
    def bind_mousewheel(self):
        """Bind mousewheel scrolling"""
        def _on_mousewheel(event):
            self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        def _bind_to_mousewheel(event):
            self.canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        def _unbind_from_mousewheel(event):
            self.canvas.unbind_all("<MouseWheel>")
        
        self.canvas.bind('<Enter>', _bind_to_mousewheel)
        self.canvas.bind('<Leave>', _unbind_from_mousewheel)

# ------------------ Enhanced GUI Application ------------------

class ModernInsuranceGUI:
    """Modern, colorful, and fully scrollable GUI application"""

    def __init__(self, root):
        self.root = root
        self.root.title("🏥 Enhanced Insurance Document Analyzer")
        self.root.geometry("1600x1000")
        self.root.configure(bg=COLORS['light'])
        
        # Initialize system components
        self.document_processor = EnhancedDocumentProcessor()
        self.query_processor = EnhancedQueryProcessor()
        self.selected_files = []
        self.system_ready = False
        self.current_result = None
        
        # Configure styles
        self.setup_styles()
        
        # Create GUI
        self.create_widgets()
        self.update_status("Ready. Upload documents to begin intelligent analysis.")

    def setup_styles(self):
        """Setup custom styles for ttk widgets"""
        style = ttk.Style()
        style.theme_use('clam')
        
        # Configure custom styles
        style.configure('Primary.TButton', 
                       background=COLORS['primary'], 
                       foreground='white',
                       font=('Segoe UI', 10, 'bold'),
                       borderwidth=0,
                       relief='flat')
        
        style.configure('Success.TButton', 
                       background=COLORS['success'], 
                       foreground='white',
                       font=('Segoe UI', 10, 'bold'),
                       borderwidth=0,
                       relief='flat')
        
        style.configure('Danger.TButton', 
                       background=COLORS['danger'], 
                       foreground='white',
                       font=('Segoe UI', 10, 'bold'),
                       borderwidth=0,
                       relief='flat')
        
        style.configure('Info.TButton', 
                       background=COLORS['info'], 
                       foreground='white',
                       font=('Segoe UI', 10, 'bold'),
                       borderwidth=0,
                       relief='flat')
        
        # Notebook styles
        style.configure('Modern.TNotebook', background=COLORS['light'])
        style.configure('Modern.TNotebook.Tab', 
                       background=COLORS['bg_primary'],
                       foreground=COLORS['text_dark'],
                       font=('Segoe UI', 11, 'bold'),
                       padding=[20, 10])

    def create_widgets(self):
        """Create modern GUI layout with full scrollability"""
        
        # Header section (fixed)
        self.create_header()
        
        # Main content with notebook
        self.notebook = ttk.Notebook(self.root, style='Modern.TNotebook')
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        
        # Create tabs with scrollable content
        self.create_upload_tab()
        self.create_query_tab()
        self.create_settings_tab()
        
        # Status bar (fixed)
        self.create_status_bar()

    def create_header(self):
        """Create colorful header section"""
        header_frame = tk.Frame(self.root, bg=COLORS['dark'], height=100)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Gradient-like effect with multiple frames
        gradient_frame = tk.Frame(header_frame, bg=COLORS['primary'], height=5)
        gradient_frame.pack(fill=tk.X)
        
        # Main header content
        content_frame = tk.Frame(header_frame, bg=COLORS['dark'])
        content_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=15)
        
        # Title with icon
        title_label = tk.Label(
            content_frame,
            text="🏥 Enhanced Insurance Document Analyzer",
            font=('Segoe UI', 24, 'bold'),
            fg='white',
            bg=COLORS['dark']
        )
        title_label.pack(side=tk.LEFT)
        
        # Status indicator
        self.system_status_label = tk.Label(
            content_frame,
            text="🔴 System Not Ready",
            font=('Segoe UI', 12, 'bold'),
            fg=COLORS['danger'],
            bg=COLORS['dark']
        )
        self.system_status_label.pack(side=tk.RIGHT)
        
        # Subtitle
        subtitle_frame = tk.Frame(header_frame, bg=COLORS['dark'])
        subtitle_frame.pack(fill=tk.X, padx=20, pady=(0, 10))
        
        subtitle_label = tk.Label(
            subtitle_frame,
            text="Powered by Google Gemini AI • Free Tier Optimized • Multi-format Document Support",
            font=('Segoe UI', 11),
            fg=COLORS['info'],
            bg=COLORS['dark']
        )
        subtitle_label.pack()

    def create_upload_tab(self):
        """Create enhanced upload tab with full scrollability"""
        # Create scrollable frame for the entire tab
        self.upload_scroll = ScrollableFrame(self.notebook, bg=COLORS['light'])
        self.notebook.add(self.upload_scroll, text="📄 Document Upload & Processing")
        
        # Container for all content
        container = self.upload_scroll.scrollable_frame
        
        # Upload section
        upload_card = self.create_card(container, "📁 Upload Documents", COLORS['bg_primary'])
        upload_card.pack(fill=tk.X, padx=20, pady=20)
        
        # Info section with colorful background
        info_frame = tk.Frame(upload_card, bg=COLORS['info'], relief=tk.RIDGE, bd=2)
        info_frame.pack(fill=tk.X, padx=15, pady=10)
        
        info_text = """
📋 Supported Document Formats:
• PDF Documents (.pdf) - Insurance policies, contracts, terms & conditions
• Word Documents (.docx) - Policy documents, claim forms, agreements  
• Images (.jpg, .jpeg, .png) - Scanned documents, policy cards, certificates
• Text Files (.txt, .md) - Plain text documents, terms & conditions

💡 Tips for Best Results:
• Ensure documents are clear and readable
• Upload all related policy documents together
• Images should be high quality for better text extraction
        """
        
        info_label = tk.Label(
            info_frame,
            text=info_text.strip(),
            font=('Segoe UI', 10),
            fg='white',
            bg=COLORS['info'],
            justify=tk.LEFT
        )
        info_label.pack(padx=15, pady=15)
        
        # Upload buttons with colorful design
        button_frame = tk.Frame(upload_card, bg=COLORS['white'])
        button_frame.pack(fill=tk.X, padx=15, pady=15)
        
        buttons = [
            ("📄 PDF Files", lambda: self.select_files([("PDF files", "*.pdf")]), COLORS['primary']),
            ("📝 Word Files", lambda: self.select_files([("Word files", "*.docx")]), COLORS['success']),
            ("🖼️ Images", lambda: self.select_files([("Image files", "*.jpg;*.jpeg;*.png")]), COLORS['warning']),
            ("📋 Text Files", lambda: self.select_files([("Text files", "*.txt;*.md")]), COLORS['secondary']),
            ("🗂️ All Files", lambda: self.select_files([("All supported", "*.*")]), COLORS['dark'])
        ]
        
        for i, (text, command, color) in enumerate(buttons):
            btn = self.create_colored_button(button_frame, text, command, color)
            btn.grid(row=i//3, column=i%3, padx=5, pady=5, sticky='ew')
        
        # Configure grid weights
        for i in range(3):
            button_frame.grid_columnconfigure(i, weight=1)
        
        # File list with enhanced display
        list_card = self.create_card(container, "📋 Selected Documents", COLORS['bg_secondary'])
        list_card.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Treeview with better styling
        tree_frame = tk.Frame(list_card, bg=COLORS['white'])
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        columns = ('Name', 'Type', 'Size', 'Status')
        self.file_tree = ttk.Treeview(tree_frame, columns=columns, show='headings', height=12)
        
        # Configure columns with colors
        for col in columns:
            self.file_tree.heading(col, text=col)
        
        self.file_tree.column('Name', width=400)
        self.file_tree.column('Type', width=100)
        self.file_tree.column('Size', width=100)
        self.file_tree.column('Status', width=200)
        
        # Scrollbars for treeview
        v_scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.file_tree.yview)
        h_scroll = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL, command=self.file_tree.xview)
        self.file_tree.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        
        self.file_tree.grid(row=0, column=0, sticky='nsew')
        v_scroll.grid(row=0, column=1, sticky='ns')
        h_scroll.grid(row=1, column=0, sticky='ew')
        
        tree_frame.grid_rowconfigure(0, weight=1)
        tree_frame.grid_columnconfigure(0, weight=1)
        
        # Control buttons with enhanced styling
        control_frame = tk.Frame(list_card, bg=COLORS['white'])
        control_frame.pack(fill=tk.X, padx=15, pady=(0, 15))
        
        self.process_btn = self.create_colored_button(
            control_frame, "🚀 Process & Analyze Documents", 
            self.process_documents, COLORS['success']
        )
        self.process_btn.pack(side=tk.LEFT, padx=5)
        
        clear_btn = self.create_colored_button(
            control_frame, "🗑️ Clear All", 
            self.clear_file_list, COLORS['danger']
        )
        clear_btn.pack(side=tk.LEFT, padx=5)
        
        # Enhanced progress bar
        progress_frame = tk.Frame(control_frame, bg=COLORS['primary'], relief=tk.RIDGE, bd=2)
        progress_frame.pack(side=tk.RIGHT, padx=10)
        
        tk.Label(progress_frame, text="Progress:", font=('Segoe UI', 10, 'bold'), 
                fg='white', bg=COLORS['primary']).pack(side=tk.LEFT, padx=5)
        
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            progress_frame, variable=self.progress_var,
            mode='determinate', length=200
        )
        self.progress_bar.pack(side=tk.LEFT, padx=5, pady=5)
        
        # Processing summary with colorful display
        summary_card = self.create_card(container, "📊 Processing Summary", COLORS['bg_primary'])
        summary_card.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        self.summary_text = scrolledtext.ScrolledText(
            summary_card, height=8, font=('Consolas', 10), 
            wrap=tk.WORD, bg=COLORS['light']
        )
        self.summary_text.pack(fill=tk.X, padx=15, pady=15)

    def create_query_tab(self):
        """Create enhanced query tab with larger answer area"""
        # Create scrollable frame
        self.query_scroll = ScrollableFrame(self.notebook, bg=COLORS['light'])
        self.notebook.add(self.query_scroll, text="🔍 Query & Analysis")
        
        container = self.query_scroll.scrollable_frame
        
        # Query input section
        query_card = self.create_card(container, "💭 Ask Questions About Your Documents", COLORS['bg_primary'])
        query_card.pack(fill=tk.X, padx=20, pady=20)
        
        # Instructions with colorful background
        instr_frame = tk.Frame(query_card, bg=COLORS['info'], relief=tk.RIDGE, bd=2)
        instr_frame.pack(fill=tk.X, padx=15, pady=10)
        
        instructions = """
💡 Ask any question about your uploaded documents:
• "What is covered under this insurance policy?"
• "What are the exclusions and waiting periods?"  
• "How do I file a claim and what documents are required?"
• "What is the sum insured amount and premium details?"
• "Are there any age or location restrictions?"
• "What are the terms and conditions for specific treatments?"
        """
        
        instr_label = tk.Label(
            instr_frame,
            text=instructions.strip(),
            font=('Segoe UI', 10),
            fg='white',
            bg=COLORS['info'],
            justify=tk.LEFT
        )
        instr_label.pack(padx=15, pady=15)
        
        # Query text area with enhanced styling
        query_input_frame = tk.Frame(query_card, bg=COLORS['white'], relief=tk.RIDGE, bd=2)
        query_input_frame.pack(fill=tk.X, padx=15, pady=10)
        
        tk.Label(query_input_frame, text="Enter your question:", 
                font=('Segoe UI', 12, 'bold'), fg=COLORS['text_dark'], 
                bg=COLORS['white']).pack(anchor=tk.W, padx=10, pady=(10, 5))
        
        self.query_text = scrolledtext.ScrolledText(
            query_input_frame, height=5, font=('Segoe UI', 12),
            wrap=tk.WORD, relief=tk.FLAT, bg=COLORS['light']
        )
        self.query_text.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        # Query buttons with colorful design
        btn_frame = tk.Frame(query_card, bg=COLORS['white'])
        btn_frame.pack(fill=tk.X, padx=15, pady=(0, 15))
        
        analyze_btn = self.create_colored_button(
            btn_frame, "🔍 Analyze Query", 
            self.process_query, COLORS['primary']
        )
        analyze_btn.pack(side=tk.LEFT, padx=5)
        
        clear_btn = self.create_colored_button(
            btn_frame, "🧹 Clear", 
            self.clear_query, COLORS['secondary']
        )
        clear_btn.pack(side=tk.LEFT, padx=5)
        
        # Sample queries with colorful buttons
        sample_card = self.create_card(container, "💡 Sample Questions", COLORS['bg_secondary'])
        sample_card.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        sample_frame = tk.Frame(sample_card, bg=COLORS['white'])
        sample_frame.pack(fill=tk.X, padx=15, pady=15)
        
        samples = [
            ("Coverage Details", "What is covered under this policy?", COLORS['primary']),
            ("Exclusions", "What are the exclusions and waiting periods?", COLORS['danger']),
            ("Claims Process", "How do I file a claim?", COLORS['success']),
            ("Premium Info", "What are the premium and sum insured details?", COLORS['warning']),
            ("Restrictions", "Are there any age restrictions?", COLORS['secondary'])
        ]
        
        for i, (title, query, color) in enumerate(samples):
            btn = self.create_colored_button(
                sample_frame, title, 
                lambda q=query: self.set_sample_query(q), color
            )
            btn.grid(row=i//3, column=i%3, padx=5, pady=5, sticky='ew')
        
        # Configure grid weights
        for i in range(3):
            sample_frame.grid_columnconfigure(i, weight=1)
        
        # Results section with MUCH LARGER area
        results_card = self.create_card(container, "📋 Analysis Results", COLORS['bg_primary'])
        results_card.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Results notebook with enhanced styling
        self.results_notebook = ttk.Notebook(results_card, style='Modern.TNotebook')
        self.results_notebook.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        
        # Answer tab with MUCH LARGER display
        answer_frame = tk.Frame(self.results_notebook, bg=COLORS['white'])
        self.results_notebook.add(answer_frame, text="💬 Answer")
        
        # Answer header
        answer_header = tk.Frame(answer_frame, bg=COLORS['success'], height=40)
        answer_header.pack(fill=tk.X)
        answer_header.pack_propagate(False)
        
        tk.Label(answer_header, text="📝 AI Analysis Result", 
                font=('Segoe UI', 14, 'bold'), fg='white', 
                bg=COLORS['success']).pack(pady=10)
        
        # Large answer display
        self.answer_text = scrolledtext.ScrolledText(
            answer_frame, font=('Segoe UI', 13), wrap=tk.WORD,
            bg=COLORS['light'], relief=tk.FLAT, borderwidth=0
        )
        self.answer_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Detailed analysis tab
        analysis_frame = tk.Frame(self.results_notebook, bg=COLORS['white'])
        self.results_notebook.add(analysis_frame, text="📚 Detailed Analysis")
        
        analysis_header = tk.Frame(analysis_frame, bg=COLORS['info'], height=40)
        analysis_header.pack(fill=tk.X)
        analysis_header.pack_propagate(False)
        
        tk.Label(analysis_header, text="🔍 Comprehensive Analysis", 
                font=('Segoe UI', 14, 'bold'), fg='white', 
                bg=COLORS['info']).pack(pady=10)
        
        self.analysis_text = scrolledtext.ScrolledText(
            analysis_frame, font=('Segoe UI', 11), wrap=tk.WORD,
            bg=COLORS['light'], relief=tk.FLAT
        )
        self.analysis_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Sources tab
        sources_frame = tk.Frame(self.results_notebook, bg=COLORS['white'])
        self.results_notebook.add(sources_frame, text="📄 Sources")
        
        sources_header = tk.Frame(sources_frame, bg=COLORS['warning'], height=40)
        sources_header.pack(fill=tk.X)
        sources_header.pack_propagate(False)
        
        tk.Label(sources_header, text="📚 Document Sources", 
                font=('Segoe UI', 14, 'bold'), fg='white', 
                bg=COLORS['warning']).pack(pady=10)
        
        # Sources treeview
        source_tree_frame = tk.Frame(sources_frame, bg=COLORS['white'])
        source_tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        source_columns = ('Source', 'Relevance', 'Content Preview')
        self.source_tree = ttk.Treeview(source_tree_frame, columns=source_columns, show='headings')
        
        for col in source_columns:
            self.source_tree.heading(col, text=col)
        
        self.source_tree.column('Source', width=200)
        self.source_tree.column('Relevance', width=100)
        self.source_tree.column('Content Preview', width=500)
        
        source_scroll = ttk.Scrollbar(source_tree_frame, orient=tk.VERTICAL, command=self.source_tree.yview)
        self.source_tree.configure(yscrollcommand=source_scroll.set)
        
        self.source_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        source_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # JSON export tab
        json_frame = tk.Frame(self.results_notebook, bg=COLORS['white'])
        self.results_notebook.add(json_frame, text="📁 Export")
        
        json_header = tk.Frame(json_frame, bg=COLORS['secondary'], height=40)
        json_header.pack(fill=tk.X)
        json_header.pack_propagate(False)
        
        tk.Label(json_header, text="💾 Export Results", 
                font=('Segoe UI', 14, 'bold'), fg='white', 
                bg=COLORS['secondary']).pack(pady=10)
        
        self.json_text = scrolledtext.ScrolledText(
            json_frame, font=('Consolas', 9), wrap=tk.WORD,
            bg=COLORS['light']
        )
        self.json_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))
        
        json_btn_frame = tk.Frame(json_frame, bg=COLORS['white'])
        json_btn_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        save_btn = self.create_colored_button(
            json_btn_frame, "💾 Save JSON", 
            self.save_json, COLORS['success']
        )
        save_btn.pack(side=tk.LEFT, padx=5)
        
        copy_btn = self.create_colored_button(
            json_btn_frame, "📋 Copy", 
            self.copy_json, COLORS['info']
        )
        copy_btn.pack(side=tk.LEFT, padx=5)

    def create_settings_tab(self):
        """Create enhanced settings tab"""
        self.settings_scroll = ScrollableFrame(self.notebook, bg=COLORS['light'])
        self.notebook.add(self.settings_scroll, text="⚙️ Settings & Status")
        
        container = self.settings_scroll.scrollable_frame
        
        # API Status section
        api_card = self.create_card(container, "📊 API Usage Status", COLORS['bg_primary'])
        api_card.pack(fill=tk.X, padx=20, pady=20)
        
        self.api_status_text = scrolledtext.ScrolledText(
            api_card, height=10, font=('Consolas', 10), 
            wrap=tk.WORD, bg=COLORS['light']
        )
        self.api_status_text.pack(fill=tk.X, padx=15, pady=15)
        
        refresh_btn = self.create_colored_button(
            api_card, "🔄 Refresh Status", 
            self.update_api_status, COLORS['info']
        )
        refresh_btn.pack(pady=(0, 15))
        
        # Configuration section
        config_card = self.create_card(container, "⚙️ System Configuration", COLORS['bg_secondary'])
        config_card.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        config_frame = tk.Frame(config_card, bg=COLORS['dark'], relief=tk.RIDGE, bd=2)
        config_frame.pack(fill=tk.X, padx=15, pady=15)
        
        config_text = f"""
🤖 Model: Gemini 1.5 Flash (Free Tier Optimized)
⚡ Rate Limit: {config.API_RATE_LIMIT} requests/minute
📄 Chunk Size: {config.CHUNK_SIZE} characters  
📊 Max Context Docs: {config.MAX_CONTEXT_DOCS}
🔄 Max Retries: {config.MAX_RETRIES}
💾 Cache Directory: {config.CACHE_DIR}
        """
        
        config_label = tk.Label(
            config_frame, text=config_text.strip(),
            font=('Consolas', 10), fg='white', bg=COLORS['dark'],
            justify=tk.LEFT
        )
        config_label.pack(padx=15, pady=15)
        
        # Cache management section
        cache_card = self.create_card(container, "🗂️ Cache Management", COLORS['bg_primary'])
        cache_card.pack(fill=tk.X, padx=20, pady=(0, 20))
        
        cache_frame = tk.Frame(cache_card, bg=COLORS['white'])
        cache_frame.pack(fill=tk.X, padx=15, pady=15)
        
        clear_cache_btn = self.create_colored_button(
            cache_frame, "🗑️ Clear Cache", 
            self.clear_cache, COLORS['warning']
        )
        clear_cache_btn.pack(side=tk.LEFT, padx=5)
        
        self.cache_info_var = tk.StringVar(value="Cache information loading...")
        cache_info_label = tk.Label(
            cache_frame, textvariable=self.cache_info_var,
            font=('Segoe UI', 10), fg=COLORS['text_dark'], bg=COLORS['white']
        )
        cache_info_label.pack(side=tk.LEFT, padx=20)
        
        self.update_cache_info()

    def create_status_bar(self):
        """Create colorful status bar"""
        status_frame = tk.Frame(self.root, bg=COLORS['dark'], height=35)
        status_frame.pack(fill=tk.X, side=tk.BOTTOM)
        status_frame.pack_propagate(False)
        
        # Status indicator
        indicator_frame = tk.Frame(status_frame, bg=COLORS['success'], width=20)
        indicator_frame.pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=5)
        
        # Status text
        self.status_var = tk.StringVar()
        self.status_bar = tk.Label(
            status_frame, textvariable=self.status_var,
            font=('Segoe UI', 10), fg='white', bg=COLORS['dark'],
            anchor=tk.W
        )
        self.status_bar.pack(fill=tk.X, padx=10, pady=5)

    # Utility methods for creating styled widgets
    
    def create_card(self, parent, title, bg_color):
        """Create a styled card widget"""
        card = tk.Frame(parent, bg=COLORS['white'], relief=tk.RAISED, bd=2)
        
        # Card header
        header = tk.Frame(card, bg=bg_color, height=50)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        title_label = tk.Label(
            header, text=title, font=('Segoe UI', 14, 'bold'),
            fg=COLORS['text_dark'], bg=bg_color
        )
        title_label.pack(pady=15)
        
        return card
    
    def create_colored_button(self, parent, text, command, color):
        """Create a colored button with hover effects"""
        btn = tk.Button(
            parent, text=text, command=command,
            font=('Segoe UI', 10, 'bold'),
            bg=color, fg='white',
            relief=tk.FLAT, bd=0,
            padx=20, pady=8,
            cursor='hand2'
        )
        
        # Hover effects
        def on_enter(e):
            btn.configure(bg=self.darken_color(color))
        
        def on_leave(e):
            btn.configure(bg=color)
        
        btn.bind("<Enter>", on_enter)
        btn.bind("<Leave>", on_leave)
        
        return btn
    
    def darken_color(self, color):
        """Darken a color for hover effect"""
        color_map = {
            COLORS['primary']: '#245a7a',
            COLORS['secondary']: '#7a2e5a',
            COLORS['success']: '#c17201',
            COLORS['danger']: '#9f3116',
            COLORS['warning']: '#c39435',
            COLORS['info']: '#47a085',
            COLORS['dark']: '#1e2a36'
        }
        return color_map.get(color, color)

    # Event handlers (keeping the same logic but with enhanced visual feedback)
    
    def select_files(self, file_types):
        """Select files with enhanced visual feedback"""
        try:
            files = filedialog.askopenfilenames(
                title="Select Documents for Analysis",
                filetypes=file_types + [("All files", "*.*")]
            )
            
            added_count = 0
            for file_path in files:
                if file_path not in [f['path'] for f in self.selected_files]:
                    try:
                        file_info = {
                            'path': file_path,
                            'name': os.path.basename(file_path),
                            'type': os.path.splitext(file_path)[1].upper(),
                            'size': self.get_file_size(file_path),
                            'status': '⏳ Ready for processing'
                        }
                        self.selected_files.append(file_info)
                        added_count += 1
                    except Exception as e:
                        logger.warning(f"Could not add file {file_path}: {e}")
            
            self.update_file_tree()
            if added_count > 0:
                self.update_status(f"✅ Added {added_count} file(s) for processing")
                
        except Exception as e:
            messagebox.showerror("Error", f"Error selecting files: {str(e)}")

    def get_file_size(self, file_path):
        """Get human readable file size"""
        try:
            size = os.path.getsize(file_path)
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size < 1024.0:
                    return f"{size:.1f} {unit}"
                size /= 1024.0
            return f"{size:.1f} TB"
        except:
            return "Unknown"

    def update_file_tree(self):
        """Update file tree display"""
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)
        
        for file_info in self.selected_files:
            self.file_tree.insert('', 'end', values=(
                file_info['name'],
                file_info['type'],
                file_info['size'],
                file_info['status']
            ))

    def clear_file_list(self):
        """Clear file list and reset system"""
        self.selected_files = []
        self.system_ready = False
        self.update_file_tree()
        self.summary_text.delete(1.0, tk.END)
        self.system_status_label.configure(text="🔴 System Not Ready", fg=COLORS['danger'])
        self.update_status("📋 File list cleared")

    def process_documents(self):
        """Process documents with enhanced visual feedback"""
        if not self.selected_files:
            messagebox.showwarning("No Files", "Please select documents first.")
            return
        
        self.process_btn.configure(state="disabled", text="🔄 Processing...")
        threading.Thread(target=self._process_documents_thread, daemon=True).start()

    def _process_documents_thread(self):
        """Process documents in separate thread"""
        try:
            total_files = len(self.selected_files)
            self.root.after(0, lambda: self.update_status("🚀 Starting document analysis..."))
            self.root.after(0, lambda: setattr(self.progress_var, 'set', 0))
            
            all_documents = []
            processed_summaries = []
            
            # Process each file
            for i, file_info in enumerate(self.selected_files):
                file_path = file_info['path']
                filename = file_info['name']
                
                try:
                    # Update status
                    self.root.after(0, lambda f=filename: self.update_status(f"📄 Processing {f}..."))
                    file_info['status'] = f"🔄 Processing..."
                    self.root.after(0, self.update_file_tree)
                    
                    # Process document
                    documents = self.document_processor.process_document(file_path)
                    all_documents.extend(documents)
                    
                    # Update success status
                    file_info['status'] = f"✅ Processed ({len(documents)} chunks)"
                    processed_summaries.append({
                        'filename': filename,
                        'chunks': len(documents),
                        'type': file_info['type']
                    })
                    
                except Exception as e:
                    logger.error(f"Error processing {filename}: {e}")
                    file_info['status'] = f"❌ Error: {str(e)[:30]}..."
                
                # Update progress
                progress = ((i + 1) / total_files) * 100
                self.root.after(0, lambda p=progress: setattr(self.progress_var, 'set', p))
                self.root.after(0, self.update_file_tree)
            
            if not all_documents:
                raise Exception("No documents were successfully processed")
            
            # Create vector store
            self.root.after(0, lambda: self.update_status("🔍 Creating intelligent document index..."))
            self.query_processor.create_vector_store(all_documents)
            
            # Update system status
            self.system_ready = True
            self.root.after(0, lambda: self.system_status_label.configure(
                text="🟢 System Ready", fg=COLORS['success']))
            
            # Show processing summary
            summary_text = self._create_processing_summary(processed_summaries, len(all_documents))
            
            self.root.after(0, lambda: self.summary_text.delete(1.0, tk.END))
            self.root.after(0, lambda: self.summary_text.insert(1.0, summary_text))
            
            self.root.after(0, lambda: self.update_status("🎉 Document analysis complete! Ready for queries"))
            self.root.after(0, lambda: messagebox.showinfo("Success",
                f"🎉 Successfully analyzed {len(processed_summaries)} documents!\n\n"
                f"📄 Created {len(all_documents)} intelligent content chunks\n"
                f"🔍 Vector search index created\n\n"
                "You can now ask questions about your documents!"))
            
        except Exception as e:
            error_msg = f"Error during document processing: {str(e)}"
            self.root.after(0, lambda: self.update_status(f"❌ {error_msg}"))
            self.root.after(0, lambda: messagebox.showerror("Processing Error", error_msg))
        finally:
            self.root.after(0, lambda: self.process_btn.configure(
                state="normal", text="🚀 Process & Analyze Documents"))

    def _create_processing_summary(self, processed_summaries, total_chunks):
        """Create colorful processing summary text"""
        summary_text = "📊 DOCUMENT ANALYSIS COMPLETE\n" + "="*70 + "\n\n"
        summary_text += f"✅ Successfully processed: {len(processed_summaries)} documents\n"
        summary_text += f"📄 Total content chunks: {total_chunks}\n"
        summary_text += f"🔍 Intelligent search index: Created\n"
        summary_text += f"🤖 AI analysis: Enabled\n\n"

        summary_text += "📋 PROCESSED DOCUMENTS:\n" + "-"*50 + "\n"
        for doc_info in processed_summaries:
            summary_text += f"• {doc_info['filename']} ({doc_info['type']}) - {doc_info['chunks']} chunks\n"

        summary_text += f"\n🎉 Ready for intelligent document querying!\n"
        summary_text += f"💡 Switch to 'Query & Analysis' tab to ask questions about your documents.\n"
        summary_text += f"🔍 The system can now understand and answer complex questions about:\n"
        summary_text += f"   • Coverage details and benefits\n"
        summary_text += f"   • Terms, conditions, and restrictions\n"
        summary_text += f"   • Claims procedures and requirements\n"
        summary_text += f"   • Waiting periods and eligibility criteria\n"

        return summary_text

    def set_sample_query(self, query):
        """Set sample query"""
        self.query_text.delete(1.0, tk.END)
        self.query_text.insert(1.0, query)

    def clear_query(self):
        """Clear query and results"""
        self.query_text.delete(1.0, tk.END)
        self.clear_results()

    def clear_results(self):
        """Clear all result displays"""
        self.answer_text.delete(1.0, tk.END)
        self.analysis_text.delete(1.0, tk.END)
        self.json_text.delete(1.0, tk.END)
        
        for item in self.source_tree.get_children():
            self.source_tree.delete(item)

    def process_query(self):
        """Process user query"""
        query = self.query_text.get(1.0, tk.END).strip()
        
        if not query:
            messagebox.showwarning("No Query", "Please enter a question.")
            return
        
        if not self.system_ready:
            messagebox.showwarning("System Not Ready", 
                "Please upload and process documents first.")
            return
        
        # Run in separate thread
        threading.Thread(target=self._process_query_thread, args=(query,), daemon=True).start()

    def _process_query_thread(self, query):
        """Process query in separate thread"""
        try:
            self.root.after(0, lambda: self.update_status("🤔 Analyzing your question..."))
            self.root.after(0, self.clear_results)
            
            # Process query
            result = self.query_processor.process_query(query)
            self.current_result = result
            
            # Display results
            self.root.after(0, lambda: self.display_results(result))
            self.root.after(0, lambda: self.update_status("✅ Query analysis complete"))
            
        except Exception as e:
            error_msg = f"Error processing query: {str(e)}"
            self.root.after(0, lambda: self.update_status(f"❌ {error_msg}"))
            self.root.after(0, lambda: messagebox.showerror("Query Error", error_msg))

    def display_results(self, result: AnalysisResponse):
        """Display comprehensive query results with enhanced formatting"""
        try:
            # Display main answer with better formatting
            answer_formatted = f"""
🎯 CONFIDENCE SCORE: {result.confidence_score:.1%}

📝 ANSWER:
{result.answer}

💡 This analysis is based on the content of your uploaded insurance documents. 
The AI has examined relevant sections and provided this response with supporting evidence.
            """.strip()
            
            self.answer_text.insert(1.0, answer_formatted)
            
            # Display detailed analysis with color-coded sections
            analysis_text = f"""
📚 COMPREHENSIVE ANALYSIS & JUSTIFICATION
{'='*80}

🔍 DETAILED JUSTIFICATION:
{'-'*40}
{result.justification}

"""
            
            if result.key_points:
                analysis_text += f"""
🔑 KEY FINDINGS:
{'-'*20}
"""
                for i, point in enumerate(result.key_points, 1):
                    analysis_text += f"{i}. {point}\n"
                analysis_text += "\n"
            
            if result.document_references:
                analysis_text += f"""
📖 SPECIFIC DOCUMENT REFERENCES:
{'-'*35}
"""
                for i, ref in enumerate(result.document_references, 1):
                    analysis_text += f"{i}. {ref}\n"
                analysis_text += "\n"
            
            if result.applicable_sections:
                analysis_text += f"""
📋 APPLICABLE POLICY SECTIONS:
{'-'*30}
"""
                for i, section in enumerate(result.applicable_sections, 1):
                    analysis_text += f"{i}. {section}\n"
            
            analysis_text += f"""

💡 ANALYSIS NOTES:
{'-'*18}
• This analysis is based on {len(result.sources)} relevant document sections
• Confidence level indicates the AI's certainty in the response
• Document references show exact quotes supporting the conclusion
• All information is extracted directly from your uploaded documents
            """
            
            self.analysis_text.insert(1.0, analysis_text)
            
            # Display sources with enhanced information
            for i, source in enumerate(result.sources):
                relevance_color = "🟢" if source.get('relevance', 0) > 0.7 else "🟡" if source.get('relevance', 0) > 0.4 else "🔴"
                self.source_tree.insert('', 'end', values=(
                    f"{relevance_color} {source.get('filename', 'Unknown')}",
                    f"{source.get('relevance', 0):.1%}",
                    source.get('content_preview', 'No preview available')
                ))
            
            # Display enhanced JSON
            json_data = result.dict()
            json_data['_metadata'] = {
                'generated_at': datetime.now().isoformat(),
                'model_used': 'gemini-1.5-flash',
                'processing_mode': 'enhanced_analysis',
                'total_sources_analyzed': len(result.sources)
            }
            
            json_str = json.dumps(json_data, indent=2, ensure_ascii=False, default=str)
            self.json_text.insert(1.0, json_str)
            
        except Exception as e:
            logger.error(f"Error displaying results: {e}")
            self.answer_text.insert(1.0, f"❌ Error displaying results: {str(e)}")

    def update_api_status(self):
        """Update API usage status display"""
        try:
            if hasattr(self.document_processor.model_client, 'usage_tracker'):
                stats = self.document_processor.model_client.get_usage_stats()
                
                status_text = f"""
📊 GEMINI API USAGE STATUS
{'='*60}

🔄 REQUEST STATISTICS:
   Total Requests: {stats['total_requests']}
   Daily Requests: {stats['daily_requests']}/1,500 (Flash tier limit)
   Requests/Minute: {stats['requests_per_minute']:.2f}/15 (Flash tier limit)

📈 USAGE ANALYSIS:
   Daily Usage: {(stats['daily_requests'] / 1500) * 100:.1f}%
   Rate Usage: {(stats['requests_per_minute'] / 15) * 100:.1f}%

🔧 ACTIVE OPTIMIZATIONS:
   ✅ Using Gemini 1.5 Flash (higher limits)
   ✅ Rate limiting: {config.API_RATE_LIMIT} requests/minute
   ✅ Response caching enabled
   ✅ Reduced token usage with smaller chunks
   ✅ Retry logic with exponential backoff
   ✅ Smart context selection

💡 RECOMMENDATIONS:
   • Monitor daily usage to stay within limits
   • Clear cache periodically to free up space
   • Use specific queries for better results
   • Upload related documents together for context
                """.strip()
                
            else:
                status_text = "📊 API status not available. Process some documents first to see usage statistics."
                
            self.api_status_text.delete(1.0, tk.END)
            self.api_status_text.insert(1.0, status_text)
            
        except Exception as e:
            logger.error(f"Error updating API status: {e}")

    def update_cache_info(self):
        """Update cache information"""
        try:
            if os.path.exists(config.CACHE_DIR):
                cache_files = [f for f in os.listdir(config.CACHE_DIR) if f.endswith('.json')]
                cache_count = len(cache_files)
                
                # Calculate cache size
                total_size = 0
                for file in cache_files:
                    try:
                        total_size += os.path.getsize(os.path.join(config.CACHE_DIR, file))
                    except:
                        pass
                
                size_mb = total_size / (1024 * 1024)
                self.cache_info_var.set(f"💾 Cache: {cache_count} responses ({size_mb:.1f} MB)")
            else:
                self.cache_info_var.set("💾 Cache: Not initialized")
        except Exception as e:
            self.cache_info_var.set(f"💾 Cache: Error - {str(e)}")

    def clear_cache(self):
        """Clear API response cache"""
        try:
            if os.path.exists(config.CACHE_DIR):
                import shutil
                shutil.rmtree(config.CACHE_DIR)
                os.makedirs(config.CACHE_DIR)
                messagebox.showinfo("Success", "✅ Cache cleared successfully!")
                self.update_cache_info()
            else:
                messagebox.showinfo("Info", "ℹ️ No cache to clear.")
        except Exception as e:
            messagebox.showerror("Error", f"❌ Failed to clear cache: {str(e)}")

    def save_json(self):
        """Save JSON results to file"""
        if not self.current_result:
            messagebox.showwarning("No Results", "No results to save.")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(self.current_result.dict(), f, indent=2, ensure_ascii=False, default=str)
                messagebox.showinfo("Success", f"✅ Results saved to {file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"❌ Failed to save: {str(e)}")

    def copy_json(self):
        """Copy JSON results to clipboard"""
        if not self.current_result:
            messagebox.showwarning("No Results", "No results to copy.")
            return
        
        try:
            json_str = json.dumps(self.current_result.dict(), indent=2, ensure_ascii=False, default=str)
            self.root.clipboard_clear()
            self.root.clipboard_append(json_str)
            messagebox.showinfo("Success", "✅ Results copied to clipboard!")
        except Exception as e:
            messagebox.showerror("Error", f"❌ Failed to copy: {str(e)}")

    def update_status(self, message):
        """Update status bar with timestamp and color coding"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        
        # Color code status messages
        if "✅" in message or "🎉" in message:
            color = COLORS['success']
        elif "❌" in message or "Error" in message:
            color = COLORS['danger']
        elif "⚠️" in message or "Warning" in message:
            color = COLORS['warning']
        elif "🔄" in message or "Processing" in message:
            color = COLORS['info']
        else:
            color = 'white'
        
        self.status_var.set(f"[{timestamp}] {message}")
        self.status_bar.configure(fg=color)

def main():
    """Main function to run the enhanced GUI application"""
    
    # Create main window
    root = tk.Tk()
    
    # Set window properties
    root.minsize(1200, 800)  # Minimum size
    
    # Try to set window icon (optional)
    try:
        root.iconbitmap('icon.ico')
    except:
        pass
    
    # Create application
    app = ModernInsuranceGUI(root)
    
    # Handle window close
    def on_closing():
        if messagebox.askokcancel("Quit", "Do you want to quit the Enhanced Insurance Document Analyzer?"):
            root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    
    # Center window on screen
    root.update_idletasks()
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")
    
    # Show enhanced welcome message
    messagebox.showinfo("Welcome", 
        "🏥 Enhanced Insurance Document Analyzer\n\n"
        "✨ NEW FEATURES:\n"
        "• Fully scrollable interface\n"
        "• Colorful, modern design\n"
        "• Much larger answer display area\n"
        "• Enhanced visual feedback\n"
        "• Improved document processing\n\n"
        "🚀 CAPABILITIES:\n"
        "• PDF, Word, Image, and Text support\n"
        "• AI-powered document analysis\n"
        "• Intelligent question answering\n"
        "• Free tier optimized\n\n"
        "📄 Please upload your insurance documents to begin!")
    
    # Start the GUI
    root.mainloop()

if __name__ == "__main__":
    main()
